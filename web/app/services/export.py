"""Render a tailored document (plain text) to PDF or DOCX (§11.8, F-10).

Both backends are pure-Python with no system libraries, so they run the same on
a laptop and on Railway. The type is deliberately plain and one-column — a first
draft the user finishes, not a designed artefact.
"""

from __future__ import annotations

import io
import re

# Markdown the model sometimes emits despite being told not to. We render plain
# text, so **bold**/#head/`code` must be stripped or they show as literal markup
# (and a leading "*" gets misread as a bullet). Keep it conservative.
_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")
_MD_BOLD_U = re.compile(r"__(.+?)__")
_MD_ITALIC = re.compile(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)")
_MD_HEAD = re.compile(r"^\s{0,3}#{1,6}\s+")
_MD_CODE = re.compile(r"`([^`]+)`")
_PHONE = re.compile(r"\+?\d[\d\s()./-]{6,}\d")


def _strip_md(line: str) -> str:
    """Fold common markdown emphasis to plain text, so headings classify right
    and no literal ``**`` / ``#`` leaks into the rendered document."""
    s = _MD_HEAD.sub("", line)
    s = _MD_BOLD.sub(r"\1", s)
    s = _MD_BOLD_U.sub(r"\1", s)
    s = _MD_CODE.sub(r"\1", s)
    s = _MD_ITALIC.sub(r"\1", s)
    return s


_BULLET_CHARS = "-•*▪◦·"
# A role/dates line: a job title followed by a date (range), whatever separator
# the model used — "Head of Strategy (Oct 2020 - Oct 2022)", "... | Nov 2022 -
# Present", "... - 2018 - 2020". We look for a separator ( ( | / or a spaced
# dash ) immediately followed by an (optional Month +) Year, or "Present".
_ROLE_TAIL = re.compile(
    r"(?:[(\|/]|\s[-–—])\s*"
    r"((?:[A-Za-z]{3,10}\.?\s+)?(?:19|20)\d{2}|[Pp]resent)\b.*$"
)
# "Company - Location" separator (hyphen/en/em dash surrounded by spaces).
_ORG_SEP = re.compile(r"\s[-–—]\s")
# For splitting an org line into company + location we also accept a comma
# ("MON.co, Kuala Lumpur"), which the dash separator above doesn't match.
_ORG_SPLIT = re.compile(r"\s[-–—]\s|,\s")


def _is_role(s: str) -> bool:
    return bool(_ROLE_TAIL.search(s)) and len(s) <= 120


def _is_heading(s: str) -> bool:
    """A section heading: short, no sentence punctuation, and either ALL CAPS,
    ends with a colon, or a short Title Case phrase. Deliberately does NOT treat
    'any capitalised line' as a heading — that over-matched company/role lines —
    nor a 'Label: value' or comma list (e.g. 'Languages: English, Italian')."""
    if len(s) > 40 or s.endswith((".", ",", ";")):
        return False
    if s.isupper() or s.endswith(":"):
        return True
    if ":" in s or "," in s:
        return False
    return s.istitle() and len(s.split()) <= 5


def parse_lines(body: str) -> list[tuple[str, str]]:
    """Parse a plain-text CV/letter into (kind, text) pairs shared by every
    renderer (PDF, DOCX, and the live HTML preview via app.js, which mirrors
    these exact rules). Kinds: name, subtitle, contact, heading, org, role,
    bullet, body, blank.

    Position matters. The first non-blank line is the name. Lines above the
    first section heading form the header: a descriptive line becomes the
    subtitle, email/phone/piped lines become contact detail. Inside the body,
    a "Company - Location" line is an org, a line with a dated parenthetical is
    a role, and short ALL-CAPS/Title lines are section headings.
    """
    raw_lines = (body or "").split("\n")
    stripped = [_strip_md(r).strip() for r in raw_lines]

    def next_is_role(i: int, depth: int) -> bool:
        """True if a role/dates line follows within the next `depth` non-blank
        lines — the signal that the current line is a company header. depth=2
        looks past a company's one-line description; depth=1 is the immediate
        line only (used for lines with no company/location separator, so a
        section heading like 'EXPERIENCE' sitting above a company isn't misread)."""
        seen = 0
        for j in range(i + 1, len(stripped)):
            if not stripped[j]:
                continue
            if _is_role(stripped[j]):
                return True
            seen += 1
            if seen >= depth:
                break
        return False

    out: list[tuple[str, str]] = []
    seen_name = False
    seen_heading = False
    subtitle_done = False
    for i, raw in enumerate(raw_lines):
        line = _strip_md(raw)
        s = line.strip()
        if not s:
            out.append(("blank", ""))
            continue
        if not seen_name:
            out.append(("name", s))
            seen_name = True
            continue
        if s[0] in _BULLET_CHARS and not (len(s) > 1 and s[1] in _BULLET_CHARS):
            out.append(("bullet", s.lstrip(_BULLET_CHARS + " \t").strip()))
            continue
        if not seen_heading:
            # Header zone. Contact detail first; then a real section heading only
            # on a STRONG signal (ALL CAPS or a trailing colon) — a short Title
            # Case line right under the name is the headline/subtitle, not a
            # section heading, so it must not close the header.
            if "@" in s or _PHONE.search(s):
                out.append(("contact", s))
                continue
            if s.isupper() or s.endswith(":"):
                seen_heading = True
                out.append(("heading", s))
                continue
            if not subtitle_done:
                out.append(("subtitle", s))
                subtitle_done = True
                continue
            out.append(("contact", s) if "|" in s else ("body", line))
            continue
        # Body zone. Role/org detection runs BEFORE the heading test so a
        # "ZALORA - Singapore" or "MON.co, Kuala Lumpur" line is an org, not a
        # mis-read heading.
        if _is_role(s):
            out.append(("role", s))
            continue
        # Company header vs. section heading. A role follows either way, so we lean
        # on strong company signals: a 'Company - Location'/'Company, City'
        # separator, or a description sentence sitting under the name. Without those,
        # a heading-shaped line ('EDUCATION', 'SKILLS') stays a heading even when a
        # role is right below it.
        short = len(s) <= 90 and not s.endswith((".", ";", ":"))
        has_sep = _ORG_SPLIT.search(s) is not None
        first_next = next((stripped[j] for j in range(i + 1, len(stripped))
                           if stripped[j]), "")
        desc_then_role = (not has_sep and next_is_role(i, 2)
                          and first_next.endswith((".", ";")))
        if has_sep:
            is_org = short and next_is_role(i, 2)
        elif desc_then_role:
            is_org = short
        else:
            is_org = short and not _is_heading(s) and next_is_role(i, 1)
        if is_org:
            out.append(("org", s))
            continue
        if _is_heading(s):
            out.append(("heading", s))
            continue
        out.append(("body", line))

    # A body line directly following a company header is that employer's one-line
    # description — tag it so renderers can italicise it, like a real CV.
    result: list[tuple[str, str]] = []
    last_real = ""
    for kind, text in out:
        if kind == "body" and last_real == "org":
            kind = "orgdesc"
        result.append((kind, text))
        if kind != "blank":
            last_real = kind
    return result


# fpdf2's core fonts are latin-1 only; tailored text can carry smart quotes,
# dashes and bullets. Fold the common ones to ASCII so the PDF stays readable
# without shipping a Unicode TTF.
_SUBS = {
    "€": "EUR ", "™": "(TM)", "®": "(R)",
    "’": "'", "‘": "'", "“": '"', "”": '"',
    "–": "-", "—": "-", "•": "-", "…": "...",
    " ": " ", "‐": "-", "‑": "-",
}


def _ascii(text: str) -> str:
    for bad, good in _SUBS.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "replace").decode("latin-1")


def to_pdf(title: str, body: str) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(20, 18, 20)

    # Reset x to the left margin + advance y after each cell; without this, fpdf2
    # leaves x at the right margin and the next multi_cell has ~0 width and raises.
    def mc(h, txt):
        pdf.multi_cell(0, h, txt, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if title:
        pdf.set_font("Helvetica", "B", 15)
        mc(8, _ascii(title))
        pdf.ln(2)

    pdf.set_font("Helvetica", "", 11)
    for line in body.split("\n"):
        line = _ascii(_strip_md(line))
        if line.strip():
            mc(6, line)
        else:
            pdf.ln(3)
    return bytes(pdf.output())


def to_docx(title: str, body: str) -> bytes:
    from docx import Document as Docx

    doc = Docx()
    if title:
        doc.add_heading(title, level=1)
    for line in body.split("\n"):
        doc.add_paragraph(_strip_md(line))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _line_kind(line: str) -> str:
    s = line.strip()
    if not s:
        return "blank"
    if s[0] in "-•*▪◦·":
        return "bullet"
    return "heading" if _is_heading(s) else "body"


_FAMILY = {"serif": "Times", "mono": "Courier", "sans": "Helvetica"}


def _split_role(text: str) -> tuple[str, str]:
    """Split a role line into (title, dates) for whatever separator was used:
    'Head of Strategy (Oct 2020 - Oct 2022)', 'Director | Nov 2022 - Present',
    'Planner - Apr 2017 - Mar 2018' -> ('...', 'Oct 2020 - Oct 2022')."""
    m = _ROLE_TAIL.search(text)
    if not m:
        return text, ""
    title = text[: m.start()].strip()
    dates = text[m.start():].strip().lstrip("(|/-–— \t").strip()
    if dates.endswith(")"):
        dates = dates[:-1].strip()
    return (title or text), dates


def _split_org(text: str) -> tuple[str, str]:
    """'ZALORA - Singapore' -> ('ZALORA', ' - Singapore'); also splits on a
    comma ('MON.co, Kuala Lumpur' -> ('MON.co', ', Kuala Lumpur'))."""
    m = _ORG_SPLIT.search(text)
    if not m:
        return text, ""
    return text[: m.start()], text[m.start():]


def to_pdf_styled(title: str, body: str, style) -> bytes:
    """Render the CV to a PDF that mirrors the candidate's own layout: name +
    subtitle + contact header with a rule, ALL-CAPS/accented section headings
    each underlined by a rule, bold company and role lines with muted dates, and
    round hanging-indent bullets. Font class, accent colour, margins, and the
    heading upper/bold treatment come from their vision-extracted style profile.

    The role title passed in is intentionally NOT printed — a tailored CV leads
    with the candidate's name, like the original."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    fam = _FAMILY.get(getattr(style, "font_class", "sans"), "Helvetica")
    accent = getattr(style, "accent_rgb", None) or (34, 34, 34)
    upper = bool(getattr(style, "heading_upper", False))
    bold_heads = getattr(style, "heading_bold", True)
    m = getattr(style, "margin_mm", 20) or 20

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(m, 15, m)

    ink = (33, 33, 33)
    muted = (110, 110, 110)
    rule = (188, 188, 188)
    hbold = "B" if bold_heads else ""
    right = pdf.w - m

    def mc(h, txt, align="L"):
        pdf.multi_cell(0, h, txt, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align=align)

    def hrule(gap_above=1.6, gap_below=2.0):
        pdf.ln(gap_above)
        y = pdf.get_y()
        pdf.set_draw_color(*rule)
        pdf.set_line_width(0.2)
        pdf.line(m, y, right, y)
        pdf.ln(gap_below)

    header_open = True          # draw the header rule once, when the body starts

    def close_header():
        nonlocal header_open
        if header_open:
            hrule(gap_above=1.2, gap_below=2.4)
            header_open = False

    for kind, text in parse_lines(body):
        text = _ascii(text)
        if kind == "name":
            pdf.set_font(fam, "B", 19)
            pdf.set_text_color(*accent)
            mc(9, text)          # keep the candidate's own casing; only sections go upper
        elif kind == "subtitle":
            pdf.set_font(fam, "", 11.5)
            pdf.set_text_color(90, 90, 90)
            mc(6, text)
        elif kind == "contact":
            pdf.set_font(fam, "", 9)
            pdf.set_text_color(*muted)
            mc(5, text)
        elif kind == "blank":
            if not header_open:
                pdf.ln(2.2)
        elif kind == "heading":
            close_header()
            pdf.ln(1.6)
            pdf.set_font(fam, hbold, 11.5)
            pdf.set_text_color(*accent)
            mc(6, text.upper() if upper else text)
            hrule(gap_above=0.6, gap_below=2.0)
        elif kind == "org":
            close_header()
            comp, loc = _split_org(text)
            pdf.ln(1.2)
            pdf.set_text_color(*ink)
            pdf.set_font(fam, "B", 11.5)
            pdf.write(5.6, comp)
            if loc:
                pdf.set_font(fam, "", 11)
                pdf.set_text_color(*muted)
                pdf.write(5.6, loc)
            pdf.ln(6)
        elif kind == "orgdesc":
            close_header()
            pdf.set_font(fam, "I", 10)
            pdf.set_text_color(90, 90, 90)
            mc(5, text)
        elif kind == "role":
            close_header()
            role, dates = _split_role(text)
            pdf.set_text_color(*ink)
            pdf.set_font(fam, "B", 11)
            pdf.write(5.4, role + (" " if dates else ""))
            if dates:
                pdf.set_font(fam, "", 10)
                pdf.set_text_color(*muted)
                pdf.write(5.4, dates)
            pdf.ln(5.6)
        elif kind == "bullet":
            close_header()
            pdf.set_font(fam, "", 10.5)
            pdf.set_text_color(*ink)
            bx, by = m + 1.4, pdf.get_y() + 2.0
            pdf.set_fill_color(*ink)
            pdf.ellipse(bx, by, 1.1, 1.1, style="F")
            pdf.set_left_margin(m + 5)
            pdf.set_x(m + 5)
            mc(5.4, text)
            pdf.set_left_margin(m)
        else:  # body
            close_header()
            pdf.set_font(fam, "", 10.5)
            pdf.set_text_color(*ink)
            mc(5.4, text)
    return bytes(pdf.output())


def to_docx_templated(orig_docx: bytes, title: str, body: str) -> bytes:
    """DOCX rendered into a copy of the user's own .docx, so its theme fonts,
    colours and named styles carry. Layout collapses to single-column."""
    from docx import Document as Docx

    doc = Docx(io.BytesIO(orig_docx))
    # Clear the existing body (paragraphs + tables) but keep styles/theme/section.
    body_el = doc.element.body
    for child in list(body_el):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            body_el.remove(child)

    def _add(text: str, style: str | None):
        try:
            return doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
        except KeyError:
            return doc.add_paragraph(text)  # style not in this template

    def _two_run(bold_part: str, rest: str, size_pt: int | None = None):
        """A paragraph whose lead is bold and remainder muted (org / role lines)."""
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        r1 = p.add_run(bold_part)
        r1.bold = True
        if rest:
            r2 = p.add_run(rest)
            r2.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)
        if size_pt:
            for r in p.runs:
                r.font.size = Pt(size_pt)
        return p

    for kind, text in parse_lines(body):
        if kind == "blank":
            _add("", None)
        elif kind == "name":
            _add(text, "Title")
        elif kind == "subtitle":
            _add(text, "Subtitle")
        elif kind == "contact":
            _add(text, "Subtitle")
        elif kind == "heading":
            _add(text, "Heading 2")
        elif kind == "org":
            comp, loc = _split_org(text)
            _two_run(comp, loc)
        elif kind == "orgdesc":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.italic = True
        elif kind == "role":
            role, dates = _split_role(text)
            _two_run(role + (" " if dates else ""), dates, size_pt=10)
        elif kind == "bullet":
            _add(text, "List Bullet")
        else:
            _add(text, None)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
